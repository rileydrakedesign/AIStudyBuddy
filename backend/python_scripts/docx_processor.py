# docx_processor.py - DOCX text extraction and conversion for document ingestion
from io import BytesIO
from typing import List, Tuple, Dict, Optional
from docx import Document
import subprocess
import tempfile
import os
import shutil
import requests

from logger_setup import log


def extract_text_from_docx(file_stream: BytesIO) -> str:
    """
    Extract all text content from a DOCX file.

    Args:
        file_stream: BytesIO stream containing DOCX file data

    Returns:
        Concatenated text from all paragraphs in the document

    Raises:
        Exception: If document cannot be parsed or is corrupted
    """
    try:
        doc = Document(file_stream)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)

        full_text = "\n\n".join(paragraphs)
        log.info(f"Extracted {len(paragraphs)} paragraphs/cells from DOCX")
        return full_text

    except Exception as e:
        log.error(f"Failed to extract text from DOCX: {e}")
        raise


def extract_docx_metadata(file_stream: BytesIO) -> Dict[str, str]:
    """
    Extract document properties (title, author) from DOCX file.

    Args:
        file_stream: BytesIO stream containing DOCX file data

    Returns:
        Dictionary with 'title' and 'author' keys

    Raises:
        Exception: If metadata cannot be extracted
    """
    try:
        doc = Document(file_stream)
        core_properties = doc.core_properties

        metadata = {
            "title": core_properties.title or "Unknown",
            "author": core_properties.author or "Unknown"
        }

        log.info(f"Extracted DOCX metadata: title='{metadata['title']}', author='{metadata['author']}'")
        return metadata

    except Exception as e:
        log.error(f"Failed to extract DOCX metadata: {e}")
        # Return defaults if metadata extraction fails
        return {"title": "Unknown", "author": "Unknown"}


def extract_docx_paragraphs(file_stream: BytesIO) -> List[Tuple[str, int]]:
    """
    Extract paragraphs with sequential numbering for page_number field compatibility.

    This function numbers paragraphs sequentially (1, 2, 3...) so they can be stored
    in the existing `page_number` field and work identically to PDF page numbers
    for citations.

    Args:
        file_stream: BytesIO stream containing DOCX file data

    Returns:
        List of tuples: (paragraph_text, paragraph_number)
        Paragraph numbers start at 1 and increment sequentially

    Raises:
        Exception: If document cannot be parsed
    """
    try:
        doc = Document(file_stream)
        paragraphs_with_numbers = []
        paragraph_num = 1

        # Extract regular paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs_with_numbers.append((text, paragraph_num))
                paragraph_num += 1

        # Extract text from tables (each cell as a separate "paragraph")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs_with_numbers.append((text, paragraph_num))
                        paragraph_num += 1

        log.info(f"Extracted {len(paragraphs_with_numbers)} numbered paragraphs from DOCX")
        return paragraphs_with_numbers

    except Exception as e:
        log.error(f"Failed to extract paragraphs from DOCX: {e}")
        raise


def get_docx_stats(file_stream: BytesIO) -> Dict[str, int]:
    """
    Get document statistics for logging and monitoring.

    Args:
        file_stream: BytesIO stream containing DOCX file data

    Returns:
        Dictionary with paragraph_count, table_count, and character_count
    """
    try:
        doc = Document(file_stream)

        paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
        table_count = len(doc.tables)

        # Count total characters
        char_count = sum(len(p.text) for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    char_count += len(cell.text)

        stats = {
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "character_count": char_count
        }

        return stats

    except Exception as e:
        log.error(f"Failed to get DOCX stats: {e}")
        return {
            "paragraph_count": 0,
            "table_count": 0,
            "character_count": 0
        }


def convert_docx_to_pdf(docx_stream: BytesIO) -> BytesIO:
    """
    Convert a DOCX file to PDF format using LibreOffice for high-fidelity conversion.
    Preserves all formatting, fonts, images, and layout from the original document.

    Args:
        docx_stream: BytesIO stream containing DOCX file data

    Returns:
        BytesIO stream containing the generated PDF

    Raises:
        Exception: If conversion fails
    """
    temp_dir = None
    try:
        log.info("[DOCX-PDF] Starting LibreOffice conversion")

        # Create temporary directory for conversion
        temp_dir = tempfile.mkdtemp(prefix='docx_conversion_')
        input_docx_path = os.path.join(temp_dir, 'input.docx')

        # Write DOCX stream to temporary file
        with open(input_docx_path, 'wb') as f:
            docx_stream.seek(0)
            f.write(docx_stream.read())

        log.info(f"[DOCX-PDF] Wrote DOCX to temp file: {input_docx_path}")

        # Try different LibreOffice binary paths (different systems have different paths)
        libreoffice_paths = [
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS
            'soffice',  # Linux/PATH
            '/usr/bin/soffice',  # Linux common
            '/opt/libreoffice/program/soffice',  # Heroku buildpack
            '/app/vendor/libreoffice/program/soffice',  # Alternative Heroku path
        ]

        soffice_cmd = None
        for path in libreoffice_paths:
            try:
                # Check if this path exists and is executable
                if os.path.exists(path):
                    soffice_cmd = path
                    log.info(f"[DOCX-PDF] Found LibreOffice at: {path}")
                    break
                # Try using it from PATH
                elif path == 'soffice':
                    result = subprocess.run(['which', 'soffice'],
                                          capture_output=True, text=True, timeout=5)
                    if result.returncode == 0:
                        soffice_cmd = 'soffice'
                        log.info(f"[DOCX-PDF] Found LibreOffice in PATH")
                        break
            except Exception as e:
                log.debug(f"[DOCX-PDF] Path {path} not available: {e}")
                continue

        if not soffice_cmd:
            raise Exception("LibreOffice not found. Please install LibreOffice or configure LIBREOFFICE_PATH environment variable.")

        # Run LibreOffice conversion
        # --headless: run without GUI
        # --convert-to pdf: convert to PDF format
        # --outdir: output directory
        cmd = [
            soffice_cmd,
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', temp_dir,
            input_docx_path
        ]

        log.info(f"[DOCX-PDF] Running command: {' '.join(cmd)}")

        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=60  # 60 second timeout
        )

        if result.returncode != 0:
            log.error(f"[DOCX-PDF] LibreOffice conversion failed: {result.stderr}")
            raise Exception(f"LibreOffice conversion failed with code {result.returncode}: {result.stderr}")

        # Read the generated PDF
        output_pdf_path = os.path.join(temp_dir, 'input.pdf')
        if not os.path.exists(output_pdf_path):
            raise Exception(f"PDF output file not found at {output_pdf_path}")

        with open(output_pdf_path, 'rb') as f:
            pdf_data = f.read()

        log.info(f"[DOCX-PDF] Successfully converted DOCX to PDF ({len(pdf_data)} bytes)")

        pdf_buffer = BytesIO(pdf_data)
        pdf_buffer.seek(0)
        return pdf_buffer

    except subprocess.TimeoutExpired:
        log.error("[DOCX-PDF] LibreOffice conversion timed out after 60 seconds")
        raise Exception("DOCX to PDF conversion timed out")
    except Exception as e:
        log.error(f"[DOCX-PDF] Failed to convert DOCX to PDF: {e}", exc_info=True)
        raise
    finally:
        # Clean up temporary directory
        if temp_dir and os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
                log.debug(f"[DOCX-PDF] Cleaned up temp directory: {temp_dir}")
            except Exception as e:
                log.warning(f"[DOCX-PDF] Failed to clean up temp directory: {e}")


def convert_docx_to_pdf_cloudmersive(docx_stream: BytesIO, api_key: str) -> BytesIO:
    """
    Convert a DOCX file to PDF using the Cloudmersive API.
    Provides high-fidelity conversion without local dependencies.

    Args:
        docx_stream: BytesIO stream containing DOCX file data
        api_key: Cloudmersive API key

    Returns:
        BytesIO stream containing the generated PDF

    Raises:
        Exception: If conversion fails or API returns error
    """
    try:
        log.info("[CLOUDMERSIVE] Starting DOCX to PDF conversion")

        # Cloudmersive conversion endpoint
        url = "https://api.cloudmersive.com/convert/docx/to/pdf"

        headers = {
            "Apikey": api_key
        }

        # Prepare file for upload
        docx_stream.seek(0)
        files = {
            "inputFile": ("document.docx", docx_stream, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }

        log.info("[CLOUDMERSIVE] Sending conversion request to API")

        # Make API request with timeout
        response = requests.post(
            url,
            headers=headers,
            files=files,
            timeout=60  # 60 second timeout
        )

        # Check response status
        if response.status_code == 200:
            pdf_data = response.content
            log.info(f"[CLOUDMERSIVE] Successfully converted DOCX to PDF ({len(pdf_data)} bytes)")

            pdf_buffer = BytesIO(pdf_data)
            pdf_buffer.seek(0)
            return pdf_buffer

        elif response.status_code == 401:
            log.error("[CLOUDMERSIVE] Authentication failed - check API key")
            raise Exception("Cloudmersive API authentication failed. Please verify your API key.")

        elif response.status_code == 429:
            log.error("[CLOUDMERSIVE] Rate limit exceeded")
            raise Exception("Cloudmersive API rate limit exceeded. Please try again later or upgrade your plan.")

        else:
            error_msg = f"Cloudmersive API returned status {response.status_code}"
            try:
                error_detail = response.json()
                error_msg += f": {error_detail}"
            except:
                error_msg += f": {response.text[:200]}"

            log.error(f"[CLOUDMERSIVE] {error_msg}")
            raise Exception(error_msg)

    except requests.exceptions.Timeout:
        log.error("[CLOUDMERSIVE] Conversion request timed out after 60 seconds")
        raise Exception("DOCX to PDF conversion timed out")

    except requests.exceptions.RequestException as e:
        log.error(f"[CLOUDMERSIVE] Network error during conversion: {e}", exc_info=True)
        raise Exception(f"Network error during DOCX to PDF conversion: {str(e)}")

    except Exception as e:
        log.error(f"[CLOUDMERSIVE] Failed to convert DOCX to PDF: {e}", exc_info=True)
        raise
